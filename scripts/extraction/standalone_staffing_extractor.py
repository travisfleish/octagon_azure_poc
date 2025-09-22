#!/usr/bin/env python3
"""
Standalone Staffing Plan Extractor
==================================

A focused script for extracting staffing plans from SOW documents with improved accuracy.
Addresses common issues with table structure preservation and extraction reliability.

Key improvements:
- Preserves table structure using python-docx and pdfplumber
- Uses JSON schema validation with retry logic
- Enhanced allocation normalization for various formats
- Targeted extraction focusing on staffing sections only
"""

import os
import json
import asyncio
import re
import jsonschema
import io
from datetime import datetime
from typing import Dict, List, Any, Optional, Tuple
from pathlib import Path
from dataclasses import dataclass

from openai import AsyncOpenAI
import pandas as pd

# Try to import optional dependencies
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not available. DOCX table extraction will be limited.")

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False
    print("Warning: pdfplumber not available. PDF table extraction will be limited.")

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False
    print("Warning: PyPDF2 not available. PDF text extraction will be limited.")

# Optional OCR stack for image-based tables
try:
    import pytesseract
    from pytesseract import Output as TesseractOutput
    PYTESSERACT_AVAILABLE = True
except Exception:
    PYTESSERACT_AVAILABLE = False
    print("Warning: pytesseract not available. OCR for image-based tables will be disabled.")

try:
    import pypdfium2 as pdfium
    PDFIUM_AVAILABLE = True
except Exception:
    PDFIUM_AVAILABLE = False
    print("Warning: pypdfium2 not available. PDF page rendering for OCR will be disabled.")

# Optional: PyMuPDF for embedded image extraction
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except Exception:
    PYMUPDF_AVAILABLE = False
    print("Warning: PyMuPDF not available. Embedded image extraction will be disabled.")


@dataclass
class StaffingEntry:
    """Individual staffing plan entry"""
    name: str
    role: str
    allocation: str
    workstream: Optional[str] = None
    level: Optional[str] = None
    location: Optional[str] = None
    pct_time: Optional[float] = None
    hours: Optional[float] = None
    months: Optional[float] = None
    notes: Optional[str] = None
    page_refs: Optional[List[int]] = None


@dataclass
class StaffingPlanResult:
    """Result of staffing plan extraction"""
    file_name: str
    staffing_plan_present: bool
    plan_type: str  # "table", "list", "mixed", "none"
    entries: List[StaffingEntry]
    totals: Dict[str, Any]
    raw_excerpt: str
    extraction_timestamp: str
    processing_time: float
    error: Optional[str] = None


class StandaloneStaffingExtractor:
    """Standalone extractor focused specifically on staffing plans"""
    
    def __init__(self):
        self.openai_client = None
        self.fte_yearly_hours_basis = 1800  # Standard FTE hours per year
        
        # JSON schema for staffing plan validation
        self.staffing_schema = {
            "type": "object",
            "properties": {
                "file_name": {"type": "string"},
                "staffing_plan_present": {"type": "boolean"},
                "plan_type": {"type": "string", "enum": ["table", "list", "mixed", "none"]},
                "entries": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "workstream": {"type": ["string", "null"]},
                            "title": {"type": "string"},
                            "level": {"type": ["string", "null"]},
                            "location": {"type": ["string", "null"]},
                            "pct_time": {"type": ["number", "null"]},
                            "hours": {"type": ["number", "null"]},
                            "months": {"type": ["number", "null"]},
                            "notes": {"type": ["string", "null"]},
                            "page_refs": {"type": "array", "items": {"type": "integer"}}
                        },
                        "required": ["title"]
                    }
                },
                "totals": {
                    "type": "object",
                    "properties": {
                        "hours": {"type": ["number", "null"]},
                        "fte_yearly_hours_basis": {"type": "number"}
                    }
                },
                "raw_excerpt": {"type": "string"}
            },
            "required": ["file_name", "staffing_plan_present", "plan_type", "entries", "totals", "raw_excerpt"]
        }
    
    async def initialize(self):
        """Initialize OpenAI client"""
        self.openai_client = AsyncOpenAI(
            api_key=os.getenv("AZURE_OPENAI_API_KEY"),
            base_url=f"{os.getenv('AZURE_OPENAI_ENDPOINT')}/openai/deployments/{os.getenv('AZURE_OPENAI_DEPLOYMENT')}",
            default_query={"api-version": os.getenv("AZURE_OPENAI_API_VERSION", "2024-08-01-preview")}
        )
    
    def extract_text_and_tables_from_docx(self, file_path: Path) -> Tuple[str, List[List[List[str]]]]:
        """Extract text and tables from DOCX file preserving structure"""
        if not DOCX_AVAILABLE:
            raise Exception("python-docx not available for DOCX processing")
        
        try:
            doc = Document(file_path)
            
            # Extract text from paragraphs
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text.strip())
            
            # Extract tables with structure preservation
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        row_data.append(cell_text)
                    table_data.append(row_data)
                tables.append(table_data)
            
            return "\n".join(text_parts), tables
            
        except Exception as e:
            raise Exception(f"Error extracting from DOCX {file_path}: {e}")
    
    def extract_text_and_tables_from_pdf(self, file_path: Path) -> Tuple[str, List[List[List[str]]]]:
        """Extract text and tables from PDF file preserving structure"""
        try:
            with open(file_path, 'rb') as f:
                file_data = f.read()
            
            text = ""
            tables = []
            
            # Try pdfplumber first for better table extraction
            if PDFPLUMBER_AVAILABLE:
                try:
                    import io
                    with pdfplumber.open(io.BytesIO(file_data)) as pdf:
                        for page in pdf.pages[:10]:  # Limit to first 10 pages
                            # Extract text
                            page_text = page.extract_text()
                            if page_text:
                                text += f"--- PAGE {len(tables) + 1} ---\n{page_text}\n"
                            
                            # Extract tables
                            page_tables = page.extract_tables()
                            for table in page_tables:
                                if table and len(table) > 1:  # Skip single-row tables
                                    tables.append(table)
                except Exception as e:
                    print(f"PDFplumber extraction failed: {e}")
            
            # Fallback to PyPDF2 if pdfplumber failed or not available
            if not text and PYPDF2_AVAILABLE:
                try:
                    import io
                    reader = PyPDF2.PdfReader(io.BytesIO(file_data))
                    for i, page in enumerate(reader.pages[:10]):
                        page_text = page.extract_text()
                        if page_text:
                            text += f"--- PAGE {i + 1} ---\n{page_text}\n"
                except Exception as e:
                    print(f"PyPDF2 extraction failed: {e}")
            
            if not text:
                raise Exception("Failed to extract text from PDF")
            
            return text, tables
            
        except Exception as e:
            raise Exception(f"Error extracting from PDF {file_path}: {e}")

    def _render_pdf_pages_to_images(self, file_path: Path, max_pages: int = 10):
        """Render PDF pages to PIL images using pypdfium2 if available"""
        if not PDFIUM_AVAILABLE:
            return []
        try:
            images = []
            doc = pdfium.PdfDocument(str(file_path))
            page_count = min(len(doc), max_pages)
            for i in range(page_count):
                page = doc[i]
                # Render at higher DPI (approx 400)
                bmp = page.render(scale=600/72)
                img = bmp.to_pil()
                # Preprocess: convert to grayscale and enhance contrast
                try:
                    from PIL import ImageOps, ImageFilter
                    g = img.convert('L')
                    g = ImageOps.autocontrast(g)
                    # Light sharpen and slight upsample for OCR clarity
                    g = g.filter(ImageFilter.SHARPEN)
                    g = g.resize((int(g.width*1.3), int(g.height*1.3)))
                    images.append(g)
                except Exception:
                    images.append(img)
            return images
        except Exception:
            return []

    def _crop_below_anchor(self, image, anchor_regex: str = r"insert\s+screen\s+shot|resource\s+table|staff\s*plan"):
        """Crop region below an anchor text to focus on embedded tables after the prompt line."""
        if not PYTESSERACT_AVAILABLE:
            return None
        try:
            data = pytesseract.image_to_data(image, output_type=TesseractOutput.DICT, config='--psm 6 -l eng')
        except Exception:
            return None
        n = len(data.get('text', []))
        min_top = None
        import re as _re
        for i in range(n):
            txt = (data['text'][i] or '').strip()
            if not txt:
                continue
            if _re.search(anchor_regex, txt, _re.I):
                top = data.get('top', [0]*n)[i]
                if min_top is None or top < min_top:
                    min_top = top
        if min_top is None:
            # Heuristic: crop lower 60%
            top = int(image.height * 0.4)
        else:
            top = int(min_top + image.height * 0.05)
        box = (0, max(0, top), image.width, image.height)
        return image.crop(box)

    def _ocr_lines_with_gaps(self, image) -> List[str]:
        """Run OCR on an image and reconstruct table-like lines by inserting column separators based on gaps."""
        if not PYTESSERACT_AVAILABLE:
            return []
        try:
            # Use tesseract with table-friendly configs
            data = pytesseract.image_to_data(
                image,
                config='--psm 6 -l eng',
                output_type=TesseractOutput.DICT
            )
        except Exception:
            return []
        n = len(data.get('text', []))
        lines: Dict[Tuple[int,int], List[Dict[str, Any]]] = {}
        for i in range(n):
            txt = (data['text'][i] or '').strip()
            if not txt:
                continue
            conf = float(data.get('conf', [0]*n)[i]) if data.get('conf') else 0.0
            if conf < 20:  # more permissive filter
                continue
            key = (data.get('block_num', [0]*n)[i], data.get('line_num', [0]*n)[i])
            item = {
                'left': data.get('left', [0]*n)[i],
                'top': data.get('top', [0]*n)[i],
                'width': data.get('width', [0]*n)[i],
                'text': txt
            }
            lines.setdefault(key, []).append(item)
        # Reconstruct per line with gap-based separators
        out_lines: List[str] = []
        for key, words in sorted(lines.items(), key=lambda kv: (kv[0][0], min(w['top'] for w in kv[1]))):
            words_sorted = sorted(words, key=lambda w: w['left'])
            if not words_sorted:
                continue
            # Compute median character width to set a dynamic gap threshold
            widths = [w['width'] for w in words_sorted if w['width']]
            avg_char = (sum(widths)/len(widths))/max(len(" ".join(w['text'] for w in words_sorted)), 1)
            gap_threshold = max(18, avg_char * 16)  # tuned heuristic
            current_left = words_sorted[0]['left'] + words_sorted[0]['width']
            parts = [words_sorted[0]['text']]
            for w in words_sorted[1:]:
                gap = w['left'] - current_left
                if gap > gap_threshold:
                    parts.append('|')
                else:
                    parts.append(' ')
                parts.append(w['text'])
                current_left = w['left'] + w['width']
            line_text = "".join(parts)
            # Keep only lines likely to be table-like or staffing-related
            if re.search(r"(name|title|role|level|hours|%|fte|allocation|rate|salary|location|billable|per\s*annum)", line_text, re.I) or re.search(r"\b\d{1,4}\b", line_text):
                out_lines.append(line_text)
        return out_lines

    def extract_tables_from_pdf_images(self, file_path: Path) -> List[str]:
        """Extract table-like markdown lines from PDF images using OCR; returns list of markdown tables as strings."""
        images = self._render_pdf_pages_to_images(file_path)
        if not images:
            return []
        markdowns: List[str] = []
        for img in images:
            # Try full page and cropped-below-anchor variants
            variants = [img]
            cropped = self._crop_below_anchor(img)
            if cropped is not None:
                variants.append(cropped)
            lines: List[str] = []
            for v in variants:
                lines = self._ocr_lines_with_gaps(v)
                if lines:
                    break
            if not lines:
                continue
            # Try to detect header line and assemble a simple markdown table
            header_idx = 0
            for i, ln in enumerate(lines[:10]):
                if self._looks_like_staffing_header(ln):
                    header_idx = i
                    break
            header = lines[header_idx] if lines else ""
            cols = max(2, header.count('|') + 1)
            sep = " | ".join(["---"] * cols)
            table_lines = [header.replace('|', ' | '), sep]
            for ln in lines[header_idx+1:]:
                table_lines.append(ln.replace('|', ' | '))
            markdowns.append("\n".join(table_lines))
        return markdowns

    def extract_tables_from_embedded_images(self, file_path: Path, save_debug: bool = True) -> List[str]:
        """Extract tables by pulling embedded images via PyMuPDF and OCR'ing them."""
        if not PYMUPDF_AVAILABLE or not PYTESSERACT_AVAILABLE:
            return []
        try:
            doc = fitz.open(str(file_path))
        except Exception:
            return []
        out: List[str] = []
        debug_dir = Path('temp/ocr_debug')
        if save_debug:
            debug_dir.mkdir(parents=True, exist_ok=True)
        import itertools
        for pno in range(min(len(doc), 10)):
            page = doc[pno]
            images = page.get_images(full=True)
            for idx, img_info in enumerate(images):
                xref = img_info[0]
                try:
                    base = doc.extract_image(xref)
                    import PIL.Image as PILImage
                    im = PILImage.open(io.BytesIO(base['image']))
                    # Preprocess
                    from PIL import ImageOps, ImageFilter
                    g = im.convert('L')
                    g = ImageOps.autocontrast(g)
                    g = g.filter(ImageFilter.SHARPEN)
                    g = g.resize((int(g.width*1.3), int(g.height*1.3)))
                    if save_debug:
                        try:
                            out_path = debug_dir / f"{file_path.stem}_p{pno+1}_img{idx+1}.png"
                            g.save(out_path)
                        except Exception:
                            pass
                    # OCR
                    lines = self._ocr_lines_with_gaps(g)
                    if not lines:
                        continue
                    # Build simple markdown table
                    header = lines[0]
                    cols = max(2, header.count('|') + 1)
                    sep = " | ".join(["---"] * cols)
                    table_lines = [header.replace('|', ' | '), sep]
                    for ln in lines[1:]:
                        table_lines.append(ln.replace('|', ' | '))
                    out.append("\n".join(table_lines))
                except Exception:
                    continue
        return out

    def parse_markdown_tables_to_entries(self, table_markdowns: List[str]) -> List[Dict[str, Any]]:
        """Best-effort deterministic parser for OCR markdown tables into staffing entries."""
        entries: List[Dict[str, Any]] = []
        for md in table_markdowns:
            lines = [ln.strip() for ln in md.split('\n') if ln.strip()]
            if len(lines) < 3:
                continue
            header = lines[0].lower()
            data_rows = [ln for ln in lines[2:] if ln and not re.search(r"total\s+sow\s+amount|total\s+hours", ln, re.I)]
            # Determine column indices
            headers = [h.strip() for h in re.split(r"\s*\|\s*", header)]
            def find_idx(preds: List[str]) -> Optional[int]:
                for i, h in enumerate(headers):
                    for p in preds:
                        if re.search(p, h, re.I):
                            return i
                return None
            idx_name = find_idx([r"name", r"personnel"]) 
            idx_title = find_idx([r"title", r"role"]) 
            idx_level = find_idx([r"level"]) 
            idx_loc = find_idx([r"location"]) 
            idx_hours = find_idx([r"billable.*hours", r"hours\b", r"#\s*hours"]) 
            idx_pct = find_idx([r"%\s*time", r"percent", r"fte\s*allocation"]) 

            for row in data_rows:
                cells = [c.strip() for c in re.split(r"\s*\|\s*", row)]
                # Skip separator lines
                if all(c == '---' for c in cells):
                    continue
                title_val = None
                name_val = None
                level_val = None
                loc_val = None
                hours_val = None
                pct_val = None
                try:
                    if idx_title is not None and idx_title < len(cells):
                        title_val = cells[idx_title]
                    if idx_name is not None and idx_name < len(cells):
                        name_val = cells[idx_name]
                    if idx_level is not None and idx_level < len(cells):
                        level_val = cells[idx_level]
                    if idx_loc is not None and idx_loc < len(cells):
                        loc_val = cells[idx_loc]
                    if idx_hours is not None and idx_hours < len(cells):
                        m = re.search(r"(\d{2,4}(?:\.\d+)?)", cells[idx_hours])
                        if m:
                            hours_val = float(m.group(1))
                    if idx_pct is not None and idx_pct < len(cells):
                        m = re.search(r"(\d{1,3}(?:\.\d+)?)\s*%", cells[idx_pct])
                        if m:
                            pct_val = float(m.group(1))
                except Exception:
                    pass

                # Fallback regex across whole row
                if not any([title_val, name_val, hours_val, pct_val]):
                    m = re.search(r"^(?P<name>[A-Za-z][A-Za-z .,'’\-]+?)\s*\|.*?(?P<title>[A-Za-z][A-Za-z /&,'’\-]+?)\s*\|.*?(?P<hours>\d{2,4}(?:\.\d+)?)\b.*?(?P<pct>\d{1,3}(?:\.\d+)?)\s*%", row)
                    if m:
                        name_val = m.group('name').strip()
                        title_val = m.group('title').strip()
                        hours_val = float(m.group('hours'))
                        pct_val = float(m.group('pct'))

                # Build entry if we have at least a title/name and one of hours/pct
                if (title_val or name_val) and (hours_val is not None or pct_val is not None):
                    title_field = title_val or ''
                    if name_val:
                        title_field = f"{name_val} / {title_field}".strip(" /")
                    entry = {
                        "title": title_field or (name_val or ""),
                        "level": level_val,
                        "location": loc_val,
                        "hours": hours_val,
                        "pct_time": pct_val,
                        "notes": None,
                        "page_refs": []
                    }
                    entries.append(entry)
        return entries
    
    def find_staffing_sections(self, text: str) -> List[str]:
        """Find sections of text that likely contain staffing information"""
        staffing_keywords = [
            'staffing', 'staff plan', 'personnel', 'team', 'resources', 'fees',
            'workstream', 'fte', 'hours', '% time', 'allocation', 'workforce',
            'resource table', 'workstream distribution', 'team members'
        ]
        
        sections = []
        lines = text.split('\n')
        
        for i, line in enumerate(lines):
            line_lower = line.lower()
            if any(keyword in line_lower for keyword in staffing_keywords):
                # Extract context around this line
                start = max(0, i - 5)
                end = min(len(lines), i + 20)
                section = '\n'.join(lines[start:end])
                sections.append(section)
        
        return sections
    
    def convert_table_to_markdown(self, table: List[List[str]]) -> str:
        """Convert table data to markdown format for better LLM processing"""
        if not table or len(table) < 2:
            return ""
        
        # Create markdown table
        markdown_lines = []
        
        # Header row
        if table[0]:
            header = " | ".join([str(cell) for cell in table[0] if cell])
            markdown_lines.append(header)
            markdown_lines.append(" | ".join(["---"] * len([cell for cell in table[0] if cell])))
        
        # Data rows
        for row in table[1:]:
            if row:
                row_text = " | ".join([str(cell) for cell in row if cell])
                if row_text.strip():
                    markdown_lines.append(row_text)
        
        return "\n".join(markdown_lines)

    def _header_tokens(self, header_line: str) -> List[str]:
        return [h.strip().lower() for h in re.split(r"\s*\|\s*", header_line)]

    def _looks_like_staffing_header(self, header_line: str) -> bool:
        headers = self._header_tokens(header_line)
        has_title = any(re.search(p, h, re.I) for h in headers for p in [r"name", r"title", r"role", r"personnel"])
        has_allocation = any(re.search(p, h, re.I) for h in headers for p in [r"hours", r"%\s*time", r"fte", r"allocation"])
        return bool(has_title and has_allocation)
    
    def normalize_allocation(self, allocation_text: str) -> Dict[str, Any]:
        """
        Enhanced allocation normalization handling various formats:
        - Hours: "45 hours" -> {"hours": 45, "pct_time": 2.5}
        - Percentage: "25%" -> {"pct_time": 25, "hours": 450}
        - FTE: "0.5 FTE" -> {"pct_time": 50, "hours": 900}
        - Months: "6 months" -> {"months": 6}
        - Mixed: "25% (450 hours)" -> {"pct_time": 25, "hours": 450}
        """
        if not allocation_text or not allocation_text.strip():
            return {"pct_time": None, "hours": None, "months": None}
        
        allocation_text = allocation_text.strip()
        result = {"pct_time": None, "hours": None, "months": None}
        
        # Extract percentage
        pct_match = re.search(r'(\d+(?:\.\d+)?)\s*%', allocation_text, re.IGNORECASE)
        if pct_match:
            result["pct_time"] = float(pct_match.group(1))
        
        # Extract hours
        hours_match = re.search(r'(\d+(?:\.\d+)?)\s*h(?:ours?|rs?)?', allocation_text, re.IGNORECASE)
        if hours_match:
            result["hours"] = float(hours_match.group(1))
        
        # Extract FTE
        fte_match = re.search(r'(\d+(?:\.\d+)?)\s*fte', allocation_text, re.IGNORECASE)
        if fte_match:
            fte_value = float(fte_match.group(1))
            result["pct_time"] = fte_value * 100
            result["hours"] = fte_value * self.fte_yearly_hours_basis
        
        # Extract months
        months_match = re.search(r'(\d+(?:\.\d+)?)\s*x\s*(\d+(?:\.\d+)?)\s*months?', allocation_text, re.IGNORECASE)
        if months_match:
            # Handle "6 x 11 months" format
            multiplier = float(months_match.group(1))
            months = float(months_match.group(2))
            result["months"] = multiplier * months
        else:
            months_simple = re.search(r'(\d+(?:\.\d+)?)\s*months?', allocation_text, re.IGNORECASE)
            if months_simple:
                result["months"] = float(months_simple.group(1))
        
        # Convert between hours and percentage if only one is present
        if result["pct_time"] and not result["hours"]:
            result["hours"] = (result["pct_time"] / 100) * self.fte_yearly_hours_basis
        elif result["hours"] and not result["pct_time"]:
            result["pct_time"] = (result["hours"] / self.fte_yearly_hours_basis) * 100
        
        return result
    
    async def extract_staffing_from_text(self, text: str, file_name: str) -> Dict[str, Any]:
        """Extract staffing plan from text using LLM with retry logic"""
        
        system_prompt = """You are a staffing plan extraction expert. Extract staffing information from SOW documents.

CRITICAL RULES:
1. Extract ONLY explicit data present in the text
2. If no staffing plan is found, set "staffing_plan_present": false
3. Preserve page numbers if provided (e.g., [p3])
4. Convert percentages to hours using 1800 hours/year FTE basis when possible
5. Return ONLY valid JSON that matches the provided schema

Look for:
- Personnel tables with names, roles, allocations
- Staffing sections with team member details
- Resource allocation tables
- Any structured data showing who works on the project and for how long

If you find staffing data, extract it into the entries array with:
- title: job title or role name
- workstream: department or workstream (if mentioned)
- level: seniority level (if mentioned)
- location: geographic location (if mentioned)
- pct_time: percentage allocation (0-100)
- hours: hours allocation
- months: duration in months
- notes: any additional notes
- page_refs: page numbers where found

Return valid JSON only."""

        user_prompt = f"""Extract staffing plan from this SOW document:

File: {file_name}

TEXT:
{text[:8000]}  # Limit to avoid token limits

Look for any staffing information including tables, lists, or structured data showing team members and their allocations."""

        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ]
        
        # Retry logic with validation
        max_retries = 3
        for attempt in range(max_retries):
            try:
                response = await self.openai_client.chat.completions.create(
                    model="gpt-5-mini",
                    messages=messages,
                    response_format={"type": "json_schema", "json_schema": {"name": "staffing_extraction", "schema": self.staffing_schema}}
                )
                
                result = json.loads(response.choices[0].message.content)
                
                # Validate against schema
                jsonschema.validate(result, self.staffing_schema)
                
                return result
                
            except (json.JSONDecodeError, jsonschema.ValidationError) as e:
                if attempt < max_retries - 1:
                    print(f"Validation failed (attempt {attempt + 1}), retrying...")
                    # Add repair prompt for next attempt
                    messages.append({
                        "role": "user", 
                        "content": f"The previous response was invalid JSON or didn't match the schema. Error: {e}. Please fix and return valid JSON."
                    })
                    continue
                else:
                    raise Exception(f"Failed to get valid JSON after {max_retries} attempts: {e}")
            except Exception as e:
                if attempt < max_retries - 1:
                    print(f"Request failed (attempt {attempt + 1}), retrying...")
                    continue
                else:
                    raise e
    
    async def extract_staffing_from_tables(self, tables: List[List[List[str]]], file_name: str) -> Dict[str, Any]:
        """Extract staffing plan from structured tables"""
        # Convert tables to markdown
        table_markdowns = []
        if tables:
            for i, table in enumerate(tables):
                markdown = self.convert_table_to_markdown(table)
                if markdown:
                    lines = [ln for ln in markdown.split('\n') if ln.strip()]
                    if lines and self._looks_like_staffing_header(lines[0]):
                        table_markdowns.append(f"Table {i + 1}:\n{markdown}")
        # Delegate to unified markdown extractor
        return await self.extract_staffing_from_table_markdowns(table_markdowns, file_name)

    async def extract_staffing_from_table_markdowns(self, table_markdowns: List[str], file_name: str) -> Dict[str, Any]:
        """Use LLM to extract staffing from a list of table markdown strings."""
        if not table_markdowns:
            return {
                "file_name": file_name,
                "staffing_plan_present": False,
                "plan_type": "none",
                "entries": [],
                "totals": {"hours": None, "fte_yearly_hours_basis": self.fte_yearly_hours_basis},
                "raw_excerpt": "No table markdowns provided"
            }

        system_prompt = """You are a staffing plan extraction expert. Extract staffing information from the provided tables.

Rules:
- Use only explicit values present in the table rows.
- Map names/titles/levels/locations.
- Convert allocations: capture hours and % when present; leave null if absent.
- Return ONLY valid JSON matching the schema."""

        user_prompt = f"File: {file_name}\n\nTABLES:\n{chr(10).join(table_markdowns[:6])}"

        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ]
        try:
            response = await self.openai_client.chat.completions.create(
                model="gpt-5-mini",
                messages=messages,
                response_format={"type": "json_schema", "json_schema": {"name": "staffing_extraction", "schema": self.staffing_schema}}
            )
            result = json.loads(response.choices[0].message.content)
            jsonschema.validate(result, self.staffing_schema)
            return result
        except Exception as e:
            print(f"Markdown table extraction failed: {e}")
            return {
                "file_name": file_name,
                "staffing_plan_present": False,
                "plan_type": "none",
                "entries": [],
                "totals": {"hours": None, "fte_yearly_hours_basis": self.fte_yearly_hours_basis},
                "raw_excerpt": f"Extraction error: {e}"
            }
    
    def convert_to_staffing_entries(self, raw_entries: List[Dict[str, Any]]) -> List[StaffingEntry]:
        """Convert raw LLM output to StaffingEntry objects with normalization"""
        entries = []
        
        for entry in raw_entries:
            # Normalize allocation
            allocation_text = entry.get("title", "") + " " + str(entry.get("pct_time", "")) + " " + str(entry.get("hours", ""))
            normalized = self.normalize_allocation(allocation_text)
            
            # Extract role and name from title
            title = entry.get("title", "")
            if " " in title:
                # Try to split title into name and role
                parts = title.split(" ", 1)
                name = parts[0] if len(parts) > 0 else "N/A"
                role = parts[1] if len(parts) > 1 else title
            else:
                name = "N/A"
                role = title
            
            staffing_entry = StaffingEntry(
                name=name,
                role=role,
                allocation=f"{normalized.get('pct_time', 0):.1f}%" if normalized.get('pct_time') else f"{normalized.get('hours', 0)} hours" if normalized.get('hours') else "N/A",
                workstream=entry.get("workstream"),
                level=entry.get("level"),
                location=entry.get("location"),
                pct_time=normalized.get("pct_time"),
                hours=normalized.get("hours"),
                months=normalized.get("months"),
                notes=entry.get("notes"),
                page_refs=entry.get("page_refs", [])
            )
            
            entries.append(staffing_entry)
        
        return entries
    
    async def extract_staffing_plan(self, file_path: Path) -> StaffingPlanResult:
        """Main method to extract staffing plan from a SOW file"""
        start_time = datetime.now()
        
        try:
            print(f"Processing {file_path.name}...")
            
            # Extract text and tables based on file type
            if file_path.suffix.lower() == '.docx':
                text, tables = self.extract_text_and_tables_from_docx(file_path)
            elif file_path.suffix.lower() == '.pdf':
                text, tables = self.extract_text_and_tables_from_pdf(file_path)
            else:
                raise Exception(f"Unsupported file type: {file_path.suffix}")
            
            # Find staffing sections in text
            staffing_sections = self.find_staffing_sections(text)
            
            # Try table extraction first if tables are available
            raw_result: Dict[str, Any] = {}
            if tables:
                print(f"Found {len(tables)} tables, attempting table extraction...")
                raw_result = await self.extract_staffing_from_tables(tables, file_path.name)

            # If no staffing found or no tables, try OCR-based table extraction for image-based PDFs
            staffing_found = bool(raw_result.get("entries")) if raw_result else False
            if (not tables or not staffing_found) and file_path.suffix.lower() == '.pdf':
                print("Attempting OCR-based table extraction from PDF images...")
                ocr_tables_markdown = self.extract_tables_from_pdf_images(file_path)
                # Also try embedded image extraction
                if not ocr_tables_markdown:
                    print("Attempting embedded image extraction via PyMuPDF...")
                    ocr_tables_markdown = self.extract_tables_from_embedded_images(file_path)
                if ocr_tables_markdown:
                    # Use LLM to extract from OCR markdowns (no deterministic parse)
                    ocr_result = await self.extract_staffing_from_table_markdowns(ocr_tables_markdown, file_path.name)
                    if not raw_result or len(ocr_result.get('entries', [])) > len(raw_result.get('entries', [])):
                        raw_result = ocr_result

            # If still nothing, fall back to text-based extraction
            if not raw_result or not raw_result.get("entries"):
                print("No staffing found from tables; attempting text extraction...")
                text_to_use = "\n".join(staffing_sections) if staffing_sections else text
                raw_result = await self.extract_staffing_from_text(text_to_use, file_path.name)
            
            # Convert to StaffingEntry objects
            entries = self.convert_to_staffing_entries(raw_result.get("entries", []))

            # Post-parse validation: drop rows missing both hours and pct_time
            entries = [e for e in entries if (e.hours is not None) or (e.pct_time is not None)]
            
            # Calculate totals
            total_hours = sum(entry.hours for entry in entries if entry.hours)
            
            processing_time = (datetime.now() - start_time).total_seconds()
            
            return StaffingPlanResult(
                file_name=file_path.name,
                staffing_plan_present=raw_result.get("staffing_plan_present", False),
                plan_type=raw_result.get("plan_type", "none"),
                entries=entries,
                totals={
                    "hours": total_hours,
                    "fte_yearly_hours_basis": self.fte_yearly_hours_basis
                },
                raw_excerpt=raw_result.get("raw_excerpt", ""),
                extraction_timestamp=datetime.utcnow().isoformat(),
                processing_time=processing_time
            )
            
        except Exception as e:
            processing_time = (datetime.now() - start_time).total_seconds()
            return StaffingPlanResult(
                file_name=file_path.name,
                staffing_plan_present=False,
                plan_type="none",
                entries=[],
                totals={"hours": None, "fte_yearly_hours_basis": self.fte_yearly_hours_basis},
                raw_excerpt=f"Error: {str(e)}",
                extraction_timestamp=datetime.utcnow().isoformat(),
                processing_time=processing_time,
                error=str(e)
            )
    
    def save_results_to_excel(self, results: List[StaffingPlanResult], filename: str = None) -> str:
        """Save staffing extraction results to Excel file"""
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"staffing_extraction_results_{timestamp}.xlsx"
        
        # Prepare data for DataFrame
        rows = []
        for result in results:
            # Create a row for each staffing entry
            if result.entries:
                for entry in result.entries:
                    rows.append({
                        "File Name": result.file_name,
                        "Staffing Present": result.staffing_plan_present,
                        "Plan Type": result.plan_type,
                        "Name": entry.name,
                        "Role": entry.role,
                        "Allocation": entry.allocation,
                        "Workstream": entry.workstream or "",
                        "Level": entry.level or "",
                        "Location": entry.location or "",
                        "Pct Time": entry.pct_time,
                        "Hours": entry.hours,
                        "Months": entry.months,
                        "Notes": entry.notes or "",
                        "Page Refs": ", ".join(map(str, entry.page_refs)) if entry.page_refs else "",
                        "Total Hours": result.totals.get("hours"),
                        "Extraction Time": result.extraction_timestamp,
                        "Processing Time": result.processing_time,
                        "Error": result.error or ""
                    })
            else:
                # No entries found
                rows.append({
                    "File Name": result.file_name,
                    "Staffing Present": result.staffing_plan_present,
                    "Plan Type": result.plan_type,
                    "Name": "",
                    "Role": "",
                    "Allocation": "",
                    "Workstream": "",
                    "Level": "",
                    "Location": "",
                    "Pct Time": None,
                    "Hours": None,
                    "Months": None,
                    "Notes": "",
                    "Page Refs": "",
                    "Total Hours": result.totals.get("hours"),
                    "Extraction Time": result.extraction_timestamp,
                    "Processing Time": result.processing_time,
                    "Error": result.error or ""
                })
        
        # Create DataFrame and save to Excel
        df = pd.DataFrame(rows)
        df.to_excel(filename, index=False, sheet_name="Staffing Extraction Results")
        
        return filename


async def main():
    """Main function to run the staffing extractor"""
    # Initialize extractor
    extractor = StandaloneStaffingExtractor()
    await extractor.initialize()
    
    # Get SOW files
    sows_dir = Path("sows")
    if not sows_dir.exists():
        print(f"SOWs directory not found: {sows_dir}")
        return
    
    sow_files = [f for f in sows_dir.iterdir() if f.is_file() and f.suffix.lower() in ['.pdf', '.docx']]
    
    if not sow_files:
        print("No SOW files found in sows directory")
        return
    
    print(f"Found {len(sow_files)} SOW files to process")
    
    # Process each file
    results = []
    for i, file_path in enumerate(sow_files, 1):
        print(f"\nProcessing {i}/{len(sow_files)}: {file_path.name}")
        result = await extractor.extract_staffing_plan(file_path)
        results.append(result)
        
        # Print summary
        if result.staffing_plan_present:
            print(f"  ✓ Found {len(result.entries)} staffing entries")
            for entry in result.entries[:3]:  # Show first 3 entries
                print(f"    - {entry.name} ({entry.role}): {entry.allocation}")
            if len(result.entries) > 3:
                print(f"    ... and {len(result.entries) - 3} more")
        else:
            print(f"  ✗ No staffing plan found")
            if result.error:
                print(f"    Error: {result.error}")
    
    # Save results
    excel_file = extractor.save_results_to_excel(results)
    print(f"\nResults saved to: {excel_file}")
    
    # Print summary
    successful = sum(1 for r in results if r.staffing_plan_present)
    total_entries = sum(len(r.entries) for r in results)
    print(f"\nSummary:")
    print(f"  Files processed: {len(results)}")
    print(f"  Files with staffing plans: {successful}")
    print(f"  Total staffing entries: {total_entries}")


if __name__ == "__main__":
    asyncio.run(main())
