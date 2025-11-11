#!/usr/bin/env python3
"""
Generate a Word document containing BOTH the general SOW extraction prompts
and the staffing plan extraction prompts (text and tables).
"""

from pathlib import Path
import re
from typing import List, Tuple
from docx import Document
from docx.shared import Pt


def add_heading(doc: Document, text: str, level: int = 1):
    doc.add_heading(text, level=level)


def add_mono_block(doc: Document, text: str):
    for line in text.splitlines():
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(10)


def find_prompt_pairs(file_text: str) -> List[Tuple[str, str]]:
    """
    Return list of (system_prompt, user_prompt) pairs in file order.
    Heuristic: For each system_prompt block, pair it with the next user_prompt block.
    """
    sys_iter = list(re.finditer(r'system_prompt\s*=\s*f?"""(.*?)"""', file_text, flags=re.DOTALL))
    usr_iter = list(re.finditer(r'user_prompt\s*=\s*f?"""(.*?)"""', file_text, flags=re.DOTALL))
    pairs: List[Tuple[str, str]] = []
    usr_idx = 0
    for sm in sys_iter:
        sys_content = sm.group(1)
        # find the next user match that occurs after this system match
        while usr_idx < len(usr_iter) and usr_iter[usr_idx].start() < sm.end():
            usr_idx += 1
        if usr_idx < len(usr_iter):
            usr_content = usr_iter[usr_idx].group(1)
            pairs.append((sys_content, usr_content))
            usr_idx += 1
        else:
            pairs.append((sys_content, ""))  # no following user prompt found
    return pairs


def main():
    project_root = Path(__file__).resolve().parents[2]
    out_dir = project_root / "docs"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "AI_Prompts_All_Extraction.docx"

    # Source files
    sow_extraction_service = project_root / "streamlit_app" / "services" / "sow_extraction_service.py"
    sow_data_extractor = project_root / "scripts" / "extraction" / "sow_data_extractor.py"
    standalone_staffing = project_root / "scripts" / "extraction" / "standalone_staffing_extractor.py"

    doc = Document()
    add_heading(doc, "AI Prompts – General SOW Extraction and Staffing Plan Extraction", level=0)

    doc.add_paragraph(
        "This document consolidates the exact system and user prompts used for SOW extraction "
        "and staffing plan extraction across the codebase. Content is pulled verbatim from source files."
    )

    # Section: General SOW Extraction
    add_heading(doc, "1. General SOW Extraction Prompts", level=2)

    # 1.a Streamlit service
    if sow_extraction_service.exists():
        text = sow_extraction_service.read_text(encoding="utf-8")
        pairs = find_prompt_pairs(text)
        add_heading(doc, "1.a SOWExtractionService (Streamlit)", level=3)
        doc.add_paragraph(str(sow_extraction_service))
        if pairs:
            # The first pair in this file corresponds to the main extraction flow
            sys_p, usr_p = pairs[0]
            add_heading(doc, "system_prompt", level=4)
            add_mono_block(doc, sys_p)
            if usr_p.strip():
                add_heading(doc, "user_prompt", level=4)
                add_mono_block(doc, usr_p)
        else:
            doc.add_paragraph("No prompts found.")

    # 1.b CLI extractor
    if sow_data_extractor.exists():
        text = sow_data_extractor.read_text(encoding="utf-8")
        pairs = find_prompt_pairs(text)
        add_heading(doc, "1.b SOW Data Extractor (Scripts)", level=3)
        doc.add_paragraph(str(sow_data_extractor))
        if pairs:
            sys_p, usr_p = pairs[0]
            add_heading(doc, "system_prompt", level=4)
            add_mono_block(doc, sys_p)
            if usr_p.strip():
                add_heading(doc, "user_prompt", level=4)
                add_mono_block(doc, usr_p)
        else:
            doc.add_paragraph("No prompts found.")

    # Section: Staffing Plan Extraction
    add_heading(doc, "2. Staffing Plan Extraction Prompts", level=2)

    if standalone_staffing.exists():
        text = standalone_staffing.read_text(encoding="utf-8")
        pairs = find_prompt_pairs(text)
        add_heading(doc, "2.a Standalone Staffing Extractor (Text)", level=3)
        doc.add_paragraph(str(standalone_staffing))
        # Heuristic: the first occurrence in this file is the text-based extraction
        if pairs:
            sys_p, usr_p = pairs[0]
            add_heading(doc, "system_prompt", level=4)
            add_mono_block(doc, sys_p)
            if usr_p.strip():
                add_heading(doc, "user_prompt", level=4)
                add_mono_block(doc, usr_p)
        else:
            doc.add_paragraph("No prompts found.")

        # Heuristic: the second occurrence is the tables-based extraction (if present)
        if len(pairs) > 1:
            add_heading(doc, "2.b Standalone Staffing Extractor (Tables)", level=3)
            sys_p2, usr_p2 = pairs[1]
            add_heading(doc, "system_prompt", level=4)
            add_mono_block(doc, sys_p2)
            if usr_p2.strip():
                add_heading(doc, "user_prompt", level=4)
                add_mono_block(doc, usr_p2)

    doc.save(str(out_path))
    print(f"✅ Wrote combined prompts document to: {out_path}")


if __name__ == "__main__":
    main()


