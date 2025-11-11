#!/usr/bin/env python3
"""
Generate a Word document containing the exact AI prompts used in extraction.
"""

from pathlib import Path
import re
from docx import Document
from docx.shared import Pt


def add_heading(doc: Document, text: str, level: int = 1):
    doc.add_heading(text, level=level)


def add_mono_block(doc: Document, text: str):
    for line in text.splitlines():
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(10)


def extract_prompt_blocks(file_path: Path, anchor_contains: str = None):
    """
    Extract all system_prompt triple-quoted strings from the given file.
    If anchor_contains is provided, only return prompts whose content includes the anchor text.
    """
    text = file_path.read_text(encoding="utf-8")
    # Capture system_prompt = """ ... """
    pattern = r'system_prompt\s*=\s*"""(.*?)"""'
    matches = re.findall(pattern, text, flags=re.DOTALL)
    if anchor_contains:
        matches = [m for m in matches if anchor_contains in m]
    return matches


def main():
    project_root = Path(__file__).resolve().parents[2]
    out_dir = project_root / "docs"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "AI_Prompts_SOW_Extraction.docx"

    # Source files
    sow_extraction_service = project_root / "streamlit_app" / "services" / "sow_extraction_service.py"
    sow_data_extractor = project_root / "scripts" / "extraction" / "sow_data_extractor.py"
    standalone_staffing = project_root / "scripts" / "extraction" / "standalone_staffing_extractor.py"

    doc = Document()
    add_heading(doc, "AI Prompts Used for SOW Extraction", level=0)

    add_heading(doc, "Overview", level=2)
    doc.add_paragraph(
        "This document contains the exact system prompts used in the codebase for SOW extraction "
        "and staffing plan extraction. They are copied verbatim from the source files to ensure fidelity."
    )

    # 1) Main SOW extraction prompt (Streamlit service)
    add_heading(doc, "SOWExtractionService (Streamlit)", level=2)
    doc.add_paragraph(str(sow_extraction_service), style=None)
    prompts = extract_prompt_blocks(sow_extraction_service, anchor_contains="You are an expert SOW analyst.")
    if prompts:
        add_heading(doc, "system_prompt", level=3)
        add_mono_block(doc, prompts[0])
    else:
        doc.add_paragraph("system_prompt not found.", style=None)

    # 2) CLI SOW extractor (same prompt, included for completeness)
    add_heading(doc, "SOW Data Extractor (Scripts)", level=2)
    doc.add_paragraph(str(sow_data_extractor), style=None)
    prompts = extract_prompt_blocks(sow_data_extractor, anchor_contains="You are an expert SOW analyst.")
    if prompts:
        add_heading(doc, "system_prompt", level=3)
        add_mono_block(doc, prompts[0])
    else:
        doc.add_paragraph("system_prompt not found.", style=None)

    # 3) Standalone staffing extractor (text-based)
    add_heading(doc, "Standalone Staffing Extractor – Text", level=2)
    doc.add_paragraph(str(standalone_staffing), style=None)
    prompts = extract_prompt_blocks(
        standalone_staffing,
        anchor_contains="You are a staffing plan extraction expert. Extract staffing information from SOW documents."
    )
    if prompts:
        add_heading(doc, "system_prompt", level=3)
        add_mono_block(doc, prompts[0])
    else:
        doc.add_paragraph("system_prompt not found.", style=None)

    # 4) Standalone staffing extractor (tables-based)
    add_heading(doc, "Standalone Staffing Extractor – Tables", level=2)
    prompts = extract_prompt_blocks(
        standalone_staffing,
        anchor_contains="You are a staffing plan extraction expert. Extract staffing information from the provided tables."
    )
    if prompts:
        add_heading(doc, "system_prompt", level=3)
        add_mono_block(doc, prompts[0])
    else:
        doc.add_paragraph("system_prompt not found.", style=None)

    doc.save(str(out_path))
    print(f"✅ Wrote prompts document to: {out_path}")


if __name__ == "__main__":
    main()


