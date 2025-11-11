#!/usr/bin/env python3
"""
Generate Staffing Plan Generator technical handover document (.docx).
"""

from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
import re
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_heading(doc: Document, text: str, level: int = 1):
    doc.add_heading(text, level=level)


def add_paragraph(doc: Document, text: str, bold: bool = False, italic: bool = False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_bullets(doc: Document, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc: Document, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_monospace_block(doc: Document, text: str):
    for line in text.splitlines():
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Courier New"
        r = run._element
        r.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        run.font.size = Pt(10)


def main():
    project_root = Path(__file__).resolve().parents[2]
    out_dir = project_root / "docs"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "Staffing_Plan_Generator_Technical_Map.docx"

    doc = Document()

    # Title
    add_heading(doc, "Staffing Plan Generator POC – Technical Map and Data Flow", level=0)
    add_paragraph(doc, f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    # 1. Overview
    add_heading(doc, "1. Overview", level=2)
    add_paragraph(
        doc,
        "This document summarizes the technical architecture, end‑to‑end data flow, technologies used, "
        "and implementation notes for the Staffing Plan Generator proof‑of‑concept (POC). "
        "It reflects the actual structure of the Streamlit app, including its four core functions."
    )

    # 1.a Streamlit App – Core Functions
    add_heading(doc, "1.a Streamlit App – Core Functions", level=3)
    add_bullets(
        doc,
        [
            "Upload SOW: Upload PDF/DOCX, extract structured data (including staffing) and export.",
            "Historical Analog: Retrieve similar/historically relevant SOWs to inform recommendations.",
            "Search: Explore indexed SOWs using lexical, semantic, or hybrid vector search.",
            "Staffing Plan Generator (Demo): Prototype that synthesizes a skeleton plan from analogs and heuristics.",
        ],
    )

    # 2. Architecture and Components
    add_heading(doc, "2. Architecture and Components", level=2)
    add_paragraph(doc, "Core components and their roles:", bold=True)
    add_bullets(
        doc,
        [
            "Streamlit App (`streamlit_app/app.py`): Interactive UI with four functions (Upload, Historical Analog, Search, Generator).",
            "SOW Extraction Service (`streamlit_app/services/sow_extraction_service.py`): "
            "Parses SOWs and extracts structured fields using Azure OpenAI; optionally leverages Azure Document Intelligence for layout/tables.",
            "Azure Document Intelligence Wrapper (`streamlit_app/services/document_intelligence_service.py`): "
            "High‑fidelity layout and table extraction for PDFs/DOCX.",
            "Azure Cognitive Search Services: REST clients for lexical/semantic and hybrid/vector queries "
            "(`streamlit_app/services/azure_search_service.py`, `vector_search_service.py`, `hybrid_search_service.py`).",
            "Indexing Scripts (`scripts/indexing/*.py`): Create and populate indexes like `octagon-sows-parsed` and `octagon-sows-hybrid`.",
            "Standalone/CLI extractors (`scripts/extraction/*.py`): Focused extraction flows (e.g., staffing‑specific extractor).",
        ],
    )

    add_paragraph(doc, "Key functions involved in generation:", bold=True)
    add_bullets(
        doc,
        [
            "Extraction: `SOWExtractionService.extract_sow_data(...)` with a targeted system prompt "
            "for staffing tables, normalization helpers (e.g., percentage/hour parsing on 1800‑hour FTE basis).",
            "UI Trigger (Generator Demo): In `app.py`, the “Generate Skeleton Staffing Plan” flow calls "
            "`analyze_staffing_patterns(...)` and `generate_skeleton_staffing_plan(...)`.",
            "Search Integration: Higher‑quality analogs are retrieved from Azure Search indices "
            "(lexical, semantic, or hybrid vector) to inform plan suggestions.",
        ],
    )

    # 3. End-to-End Data Flow
    add_heading(doc, "3. End‑to‑End Data Flow", level=2)
    add_paragraph(doc, "Aligned with the four Streamlit functions:", bold=True)
    add_numbered(
        doc,
        [
            "Upload SOW: The app processes PDF/DOCX via Azure Document Intelligence (layout/tables) and Azure OpenAI (extraction prompt) to produce structured JSON (client, project metadata, scope, deliverables, exclusions, staffing).",
            "Historical Analog: Uses Azure Cognitive Search (lexical/semantic/hybrid) to retrieve similar SOWs; results feed into downstream analysis.",
            "Search: Provides free‑form search across indexed SOW fields; supports filters and different search modes.",
            "Staffing Plan Generator (Demo): Combines patterns from analogs with simple heuristics to propose a skeleton staffing plan.",
            "Normalization: Staffing allocations normalized to percentages (1800‑hour/year basis) for consistency.",
            "Storage/Indexing (optional): Parsed artifacts can be persisted to Azure Storage and indexed via scripts into `octagon-sows-parsed` / `octagon-sows-hybrid`.",
            "Export/Review: Users can export results and iterate.",
        ],
    )

    # Diagram image
    diagram_path = out_dir / "staffing_flow_diagram.png"
    if diagram_path.exists():
        add_paragraph(doc, "Diagram:", bold=True)
        try:
            doc.add_picture(str(diagram_path), width=Inches(7.0))
        except Exception:
            add_paragraph(doc, f"(Could not embed image; found file at {diagram_path})", italic=True)
    else:
        add_paragraph(doc, f"(Flow diagram not found at {diagram_path})", italic=True)

    # 3.b AI Extraction Details (Prompts + Schema usage)
    add_heading(doc, "3.b AI Extraction Details", level=3)
    add_paragraph(
        doc,
        "The app uses Azure OpenAI with structured prompts and a JSON schema to extract "
        "client, project title, dates, project length, scope summary, deliverables, exclusions, and staffing. "
        "Deliverables and exclusions are extracted by the model per prompt + schema (no custom heuristics). "
        "Post‑processing is minimal (e.g., compute project length from dates, normalize staffing allocations)."
    )

    # Pull prompts and key snippets from code
    sow_service_path = project_root / "streamlit_app" / "services" / "sow_extraction_service.py"
    try:
        sow_text = sow_service_path.read_text(encoding="utf-8")
    except Exception:
        sow_text = ""

    # Extract primary system_prompt and user_prompt
    sys_match = re.search(r'system_prompt\s*=\s*"""(.*?)"""', sow_text, flags=re.DOTALL)
    usr_match = re.search(r'user_prompt\s*=\s*f?"""(.*?)"""', sow_text, flags=re.DOTALL)

    add_heading(doc, "Prompts (General SOW Extraction)", level=4)
    if sys_match:
        add_paragraph(doc, "system_prompt:", bold=True)
        add_monospace_block(doc, sys_match.group(1).strip())
    else:
        add_paragraph(doc, "(system_prompt not found in code)", italic=True)
    if usr_match:
        add_paragraph(doc, "user_prompt:", bold=True)
        add_monospace_block(doc, usr_match.group(1).strip())
    else:
        add_paragraph(doc, "(user_prompt not found in code)", italic=True)

    # Show where schema is enforced via response_format
    add_heading(doc, "Schema Enforcement via response_format", level=4)
    add_monospace_block(
        doc,
        "response = await self.openai_client.chat.completions.create(\n"
        "    model=\"gpt-5-mini\",\n"
        "    messages=messages,\n"
        "    response_format={\"type\": \"json_schema\", \"json_schema\": {\"name\": \"sow_extraction\", \"schema\": json_schema}}\n"
        ")"
    )

    # Attempt to extract the JSON schema Python dict block from code
    def _extract_schema_block(src: str):
        m = re.search(r'json_schema\s*=\s*\{', src)
        if not m:
            return None
        brace_start = src.find('{', m.end() - 1)
        if brace_start == -1:
            return None
        depth = 0
        i = brace_start
        while i < len(src):
            ch = src[i]
            if ch == '{':
                depth += 1
            elif ch == '}':
                depth -= 1
                if depth == 0:
                    return src[brace_start:i+1]
            i += 1
        return None

    schema_block = _extract_schema_block(sow_text) if sow_text else None
    add_heading(doc, "JSON Schema (as defined in code)", level=4)
    if schema_block:
        add_monospace_block(doc, schema_block.strip())
    else:
        add_paragraph(doc, "(Could not locate json_schema block in code)", italic=True)

    # Include date-length calculation snippet
    add_heading(doc, "Date → Project Length Post‑processing", level=4)
    add_monospace_block(
        doc,
        "def calculate_project_length(self, start_date_str: str, end_date_str: str) -> str:\n"
        "    if not start_date_str or not end_date_str:\n"
        "        return None\n"
        "    start_date = datetime.strptime(start_date_str, \"%Y-%m-%d\").date()\n"
        "    end_date = datetime.strptime(end_date_str, \"%Y-%m-%d\").date()\n"
        "    days = (end_date - start_date).days\n"
        "    if days >= 365:\n"
        "        months = round(days / 30.44)\n"
        "        return f\"{months} months\"\n"
        "    elif days >= 30:\n"
        "        weeks = round(days / 7)\n"
        "        return f\"{weeks} weeks\"\n"
        "    else:\n"
        "        return f\"{days} days\""
    )

    # 4. Configuration and Secrets
    add_heading(doc, "4. Configuration and Secrets", level=2)
    add_paragraph(doc, "Environment variables (via `.env`) drive connectivity:", bold=True)
    add_bullets(
        doc,
        [
            "Azure OpenAI: `AZURE_OPENAI_API_KEY`, `AZURE_OPENAI_ENDPOINT`, `AZURE_OPENAI_DEPLOYMENT`, `AZURE_OPENAI_API_VERSION`",
            "Azure Storage: `AZURE_STORAGE_CONNECTION_STRING` or `AZURE_STORAGE_ACCOUNT_URL` (+ `AZURE_STORAGE_KEY` or SAS)",
            "Azure Cognitive Search: `SEARCH_ENDPOINT`, `SEARCH_KEY`",
            "Azure Document Intelligence: `AZURE_DOCINT_ENDPOINT`, `AZURE_DOCINT_API_KEY`, `AZURE_DOCINT_API_VERSION`",
        ],
    )
    add_paragraph(
        doc,
        "Notes: The app gracefully degrades when certain services are not configured (e.g., skips uploads or search). "
        "SSL verification issues were mitigated by pointing `REQUESTS_CA_BUNDLE` and `SSL_CERT_FILE` to the `certifi` bundle."
    )

    # 5. Technologies Used
    add_heading(doc, "5. Technologies Used", level=2)
    add_bullets(
        doc,
        [
            "Python, Streamlit UI",
            "Azure OpenAI (extraction prompts, embeddings for vector/hybrid search)",
            "Azure Document Intelligence (prebuilt-layout) for layout/tables",
            "Azure Storage (Blob) for source/extracted/parsed artifacts",
            "Azure Cognitive Search (lexical, semantic, vector/hybrid indexes)",
            "Pandas for tabular exports; `requests`, `aiohttp`, `dotenv`, `certifi` for plumbing",
        ],
    )

    # 6. Implementation Notes and Issues Encountered
    add_heading(doc, "6. Implementation Notes and Issues Encountered", level=2)
    add_bullets(
        doc,
        [
            "SSL certificate verification with Azure endpoints: resolved by using `certifi` and setting `REQUESTS_CA_BUNDLE` and `SSL_CERT_FILE`.",
            "Async in Streamlit: event‑loop guard with fallback to a new loop when invoking async extraction.",
            "Staffing normalization: standardized to percentage using a 1800‑hour/year FTE basis; parsing handles '%', hours, and FTE.",
            "Variability of staffing tables: extraction prompt emphasizes table patterns; header canonicalization improves robustness.",
            "Service availability: components are optional; the app skips uploads/search when env vars are missing.",
            "Index naming: `octagon-sows-parsed` (structured JSON) and `octagon-sows-hybrid` (semantic/vector fields).",
            "Generator status: Staffing Plan Generator is a demo/prototype, illustrating what could be built next.",
        ],
    )

    # 7. Where to Start (Handover)
    add_heading(doc, "7. Where to Start", level=2)
    add_numbered(
        doc,
        [
            "Install dependencies: `pip install -r streamlit_app/requirements.txt`.",
            "Create `.env` at repo root with Azure keys for OpenAI, Storage, Search, and (optionally) Document Intelligence.",
            "Run the UI: `cd streamlit_app && streamlit run app.py`.",
            "Indexing (optional): use scripts in `scripts/indexing/` to create/populate Azure Search indexes.",
            "Try it: Use the four functions—Upload a SOW, explore Historical Analogs, run Search, then (optionally) generate a demo skeleton staffing plan.",
        ],
    )

    # Save
    doc.save(str(out_path))

    print(f"✅ Wrote handover document to: {out_path}")


if __name__ == "__main__":
    main()


